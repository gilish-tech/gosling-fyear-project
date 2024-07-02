from fastapi import FastAPI, File, UploadFile
from fastapi.responses import FileResponse
import fitz  # PyMuPDF
from docx import Document
from fpdf import FPDF
from happytransformer import TTSettings
import os
from model.model_loader import ModelLoader
import nltk
from textblob import TextBlob
from constants import MY_WORD_LIST


model = ModelLoader.get_model()
app = FastAPI()


def extract_text_from_pdf(file_path):
    document = fitz.open(file_path)
    text = ""
    for page_num in range(len(document)):
        page = document.load_page(page_num)
        text += page.get_text()
    return text

def extract_text_from_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        text += paragraph.text + "\n"
    return text

def split_into_sentences(text):
    return nltk.tokenize.sent_tokenize(text)

# correct spelling
def correct_spelling(text):
    if text.lower() in MY_WORD_LIST:
      return text
    blob = TextBlob(text)
    return blob.correct().string

def correct_word_in_sentence(sents):
  words = sents.split(" ")
  corrected_word =[correct_spelling(text) for text in words]
  return " ".join(corrected_word)

def correct_document(text):
    sentences = split_into_sentences(text)
    corrected_sentences = [correct_text(correct_word_in_sentence(sentence)) for sentence in sentences]
    return ' '.join(corrected_sentences)

def correct_text(text):
    args = TTSettings(
        min_length=len(text.split()),  # Set a minimum length based on input word count
        max_length=len(text.split()) + 50,  # Allow some flexibility with max length
        early_stopping=True
    )
    try:
        result = model.generate_text(text,args=args)
        print(result)
        print()
        print()
        print()
        print(len(text.split(" ")))
        print()
        print()
        # print(text)
        print()
        print()
        print()
        print()
        return result.text
    except:
        return text
def save_text_to_pdf(text, output_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    for line in text.split('\n'):
        pdf.cell(200, 10, txt=line, ln=True, align='L')
    pdf.output(output_path)

def save_text_to_docx(text, output_path):
    document = Document()
    for line in text.split('\n'):
        document.add_paragraph(line)
    document.save(output_path)

@app.post("/correct-pdf/")
async def correct_pdf(file: UploadFile = File(...)):
    print("pdf detected.....")
    print("pdf detected.....")
    print("pdf detected.....")
    print("pdf detected.....")
    input_path = f"temp_{file.filename}"
    output_path = f"corrected_{file.filename}"

    # Save uploaded file
    with open(input_path, "wb") as f:
        f.write(file.file.read())

    # Process the file
    text = extract_text_from_pdf(input_path)
    corrected_text = correct_text(text)
    save_text_to_pdf(corrected_text, output_path)

    # Clean up the temporary input file
    os.remove(input_path)

    # Return the corrected file
    return FileResponse(output_path, media_type="application/pdf", filename=output_path)


